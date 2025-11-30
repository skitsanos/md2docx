"""
Word document generator from Markdown AST.

This module converts a Markdown AST into a Word document using python-docx,
with support for custom branding (fonts, colors, headers, footers).
"""

from typing import Any, Optional, Union
from pathlib import Path
from io import BytesIO
import tempfile
import os

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

from .branding import BrandingConfig
from .parser import MarkdownParser


class WordGenerator:
    """
    Generate Word documents from Markdown AST with branding support.

    This class takes a parsed Markdown AST and converts it into a
    formatted Word document, applying custom branding configuration.
    """

    def __init__(self, branding: Optional[BrandingConfig] = None):
        """
        Initialize the Word generator.

        Args:
            branding: Optional branding configuration. If not provided,
                     default branding will be used.
        """
        self.branding = branding or BrandingConfig()
        self._document: Optional[Document] = None
        self._list_level = 0

    def generate(self, ast: list[dict[str, Any]]) -> Document:
        """
        Generate a Word document from a Markdown AST.

        Args:
            ast: The Markdown AST (list of nodes).

        Returns:
            A python-docx Document object.
        """
        self._document = Document()
        self._setup_document()
        self._process_nodes(ast)
        self._add_headers_footers()
        return self._document

    def generate_from_markdown(self, markdown_content: str) -> Document:
        """
        Generate a Word document directly from Markdown content.

        Args:
            markdown_content: The Markdown text content.

        Returns:
            A python-docx Document object.
        """
        parser = MarkdownParser()
        ast = parser.parse(markdown_content)
        return self.generate(ast)

    def save(self, document: Document, output_path: Union[str, Path]) -> None:
        """
        Save a Word document to file.

        Args:
            document: The Document object to save.
            output_path: Path where the document will be saved.
        """
        document.save(str(output_path))

    def to_bytes(self, document: Document) -> bytes:
        """
        Convert a Word document to bytes.

        Args:
            document: The Document object to convert.

        Returns:
            The document as bytes (for API responses).
        """
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def _setup_document(self) -> None:
        """Set up document properties and styles based on branding."""
        if not self._document:
            return

        # Set document properties
        core_props = self._document.core_properties
        if self.branding.title:
            core_props.title = self.branding.title
        if self.branding.author:
            core_props.author = self.branding.author
        if self.branding.company:
            core_props.company = self.branding.company

        # Set page margins
        for section in self._document.sections:
            section.page_width = self.branding.page_width
            section.page_height = self.branding.page_height
            section.top_margin = self.branding.margin_top
            section.bottom_margin = self.branding.margin_bottom
            section.left_margin = self.branding.margin_left
            section.right_margin = self.branding.margin_right

        # Set default paragraph style
        style = self._document.styles["Normal"]
        font = style.font
        font.name = self.branding.body_font.name
        font.size = self.branding.body_font.size
        if self.branding.body_font.color:
            font.color.rgb = self.branding.body_font.color

    def _add_headers_footers(self) -> None:
        """Add headers and footers to the document."""
        if not self._document:
            return

        for section in self._document.sections:
            # Add header if any header content is configured
            header_config = self.branding.header
            if header_config.text or header_config.left_text or header_config.right_text:
                header = section.header
                header.is_linked_to_previous = False  # Ensure explicit definition
                paragraph = header.paragraphs[0]

                # Build multi-zone header: Left\tCenter\tRight
                self._build_header_footer_content(paragraph, header_config, is_header=True)

                # Apply Header style
                paragraph.style = self._document.styles["Header"]

            # Add footer if any footer content is configured
            footer_config = self.branding.footer
            if footer_config.text or footer_config.left_text or footer_config.right_text or footer_config.include_page_number:
                footer = section.footer
                footer.is_linked_to_previous = False  # Ensure explicit definition
                paragraph = footer.paragraphs[0]

                # Build multi-zone footer: Left\tCenter\tRight
                self._build_header_footer_content(paragraph, footer_config, is_header=False)

                # Apply Footer style
                paragraph.style = self._document.styles["Footer"]

    def _build_header_footer_content(self, paragraph, config, is_header: bool = True) -> None:
        """Build header/footer content with left/center/right zones using tabs."""
        # Clear existing content
        paragraph.clear()

        # Determine what goes in each zone
        left_content = config.left_text
        center_content = config.text
        right_content = config.right_text

        # Handle page number placement
        if not is_header and config.include_page_number:
            page_num_text = "Page "
            if config.page_number_position == "left":
                left_content = page_num_text if not left_content else f"{left_content} - {page_num_text}"
            elif config.page_number_position == "center":
                center_content = page_num_text if not center_content else f"{center_content} - {page_num_text}"
            else:  # right (default)
                right_content = page_num_text if not right_content else f"{right_content} - {page_num_text}"

        # Add logo if configured
        logo_added = False
        if config.logo_path:
            logo_added = self._add_logo_to_zone(paragraph, config)

        # Build the content with tabs
        # Format: Left\tCenter\tRight
        # If logo is in left position, it's already added
        if not (logo_added and config.logo_position == "left"):
            run = paragraph.add_run(left_content)
            self._apply_header_footer_font(run, config)

        paragraph.add_run("\t")

        if not (logo_added and config.logo_position == "center"):
            run = paragraph.add_run(center_content)
            self._apply_header_footer_font(run, config)

        paragraph.add_run("\t")

        if not (logo_added and config.logo_position == "right"):
            run = paragraph.add_run(right_content)
            self._apply_header_footer_font(run, config)

        # Add page number field if needed
        if not is_header and config.include_page_number:
            self._add_page_number(paragraph)

    def _add_logo_to_zone(self, paragraph, config) -> bool:
        """
        Add logo image to the specified zone in header/footer.

        Returns True if logo was successfully added, False otherwise.
        """
        if not config.logo_path:
            return False

        try:
            from .images import load_image

            # Load the image (handles URLs and SVG conversion)
            image_stream = load_image(config.logo_path)

            # Determine logo width
            if config.logo_width:
                width = Inches(config.logo_width)
            else:
                width = Inches(0.5)  # Default width

            # Add to the correct position
            if config.logo_position == "left":
                run = paragraph.add_run()
                run.add_picture(image_stream, width=width)
                paragraph.add_run(" ")  # Space after logo
            elif config.logo_position == "center":
                paragraph.add_run("\t")  # Tab to center
                run = paragraph.add_run()
                run.add_picture(image_stream, width=width)
            else:  # right
                paragraph.add_run("\t\t")  # Tabs to right
                run = paragraph.add_run()
                run.add_picture(image_stream, width=width)

            return True

        except Exception as e:
            # Log error but don't fail the document generation
            print(f"Warning: Failed to add logo: {e}")
            return False

    def _apply_header_footer_font(self, run, config) -> None:
        """Apply font settings to a header/footer run."""
        run.font.name = config.font_name
        run.font.size = config.font_size
        if config.color:
            run.font.color.rgb = config.color

    def _add_page_number(self, paragraph) -> None:
        """Add page number field to a paragraph."""
        run = paragraph.add_run()
        fldChar = OxmlElement("w:fldChar")
        fldChar.set(qn("w:fldCharType"), "begin")
        run._r.append(fldChar)

        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = "PAGE"
        run._r.append(instrText)

        fldChar = OxmlElement("w:fldChar")
        fldChar.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar)

    def _process_nodes(self, nodes: list[dict[str, Any]]) -> None:
        """Process a list of AST nodes."""
        for node in nodes:
            self._process_node(node)

    def _process_node(self, node: dict[str, Any]) -> None:
        """Process a single AST node based on its type."""
        node_type = node.get("type", "")

        handlers = {
            "heading": self._handle_heading,
            "paragraph": self._handle_paragraph,
            "list": self._handle_list,
            "block_code": self._handle_code_block,
            "block_quote": self._handle_block_quote,
            "thematic_break": self._handle_thematic_break,
            "table": self._handle_table,
            "newline": self._handle_newline,
        }

        handler = handlers.get(node_type)
        if handler:
            handler(node)

    def _handle_heading(self, node: dict[str, Any]) -> None:
        """Handle heading nodes (h1-h6)."""
        if not self._document:
            return

        level = node.get("attrs", {}).get("level", 1)
        text = MarkdownParser.get_node_text(node)

        # Use Word's built-in heading style
        style_name = f"Heading {level}"
        paragraph = self._document.add_paragraph(text, style=style_name)

        # Override with branding config if custom settings are provided
        config = self.branding.get_heading_config(level)

        # Only apply custom formatting if branding has non-default values
        if config.color or config.font_name != "Calibri":
            for run in paragraph.runs:
                if config.font_name != "Calibri":
                    run.font.name = config.font_name
                if config.color:
                    run.font.color.rgb = config.color
                # Override bold/italic only if explicitly set different from default
                run.font.bold = config.bold
                if hasattr(config, "italic") and config.italic:
                    run.font.italic = config.italic

            # Set custom spacing if provided
            paragraph.paragraph_format.space_before = config.space_before
            paragraph.paragraph_format.space_after = config.space_after

    def _handle_paragraph(self, node: dict[str, Any]) -> None:
        """Handle paragraph nodes with inline formatting."""
        if not self._document:
            return

        paragraph = self._document.add_paragraph()
        self._process_inline_content(paragraph, node.get("children", []))

    def _process_inline_content(self, paragraph, children: list[dict[str, Any]]) -> None:
        """Process inline content (text, emphasis, strong, links, code)."""
        for child in children:
            child_type = child.get("type", "")

            if child_type == "text":
                run = paragraph.add_run(child.get("raw", ""))
                run.font.name = self.branding.body_font.name
                run.font.size = self.branding.body_font.size
                if self.branding.body_font.color:
                    run.font.color.rgb = self.branding.body_font.color

            elif child_type == "strong":
                text = MarkdownParser.get_node_text(child)
                run = paragraph.add_run(text)
                run.font.bold = True
                run.font.name = self.branding.body_font.name
                run.font.size = self.branding.body_font.size

            elif child_type == "emphasis":
                text = MarkdownParser.get_node_text(child)
                run = paragraph.add_run(text)
                run.font.italic = True
                run.font.name = self.branding.body_font.name
                run.font.size = self.branding.body_font.size

            elif child_type == "codespan":
                text = child.get("raw", "")
                run = paragraph.add_run(text)
                run.font.name = self.branding.code_font.name
                run.font.size = self.branding.code_font.size
                if self.branding.code_font.color:
                    run.font.color.rgb = self.branding.code_font.color

            elif child_type == "link":
                text = MarkdownParser.get_node_text(child)
                url = child.get("link", "")
                self._add_hyperlink(paragraph, url, text)

            elif child_type == "softbreak":
                paragraph.add_run(" ")

            elif child_type == "linebreak":
                paragraph.add_run("\n")

            elif child_type == "strikethrough":
                text = MarkdownParser.get_node_text(child)
                run = paragraph.add_run(text)
                run.font.strike = True

    def _add_hyperlink(self, paragraph, url: str, text: str) -> None:
        """Add a clickable hyperlink to a paragraph."""
        # Create the relationship
        part = paragraph.part
        r_id = part.relate_to(
            url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True
        )

        # Create the hyperlink element
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        # Create a new run for the hyperlink text
        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")

        # Set hyperlink styling
        color = OxmlElement("w:color")
        color.set(qn("w:val"), f"{self.branding.link_color[0]:02X}{self.branding.link_color[1]:02X}{self.branding.link_color[2]:02X}")
        rPr.append(color)

        if self.branding.link_underline:
            underline = OxmlElement("w:u")
            underline.set(qn("w:val"), "single")
            rPr.append(underline)

        # Set font
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), self.branding.body_font.name)
        rFonts.set(qn("w:hAnsi"), self.branding.body_font.name)
        rPr.append(rFonts)

        new_run.append(rPr)

        # Add the text
        text_elem = OxmlElement("w:t")
        text_elem.text = text
        new_run.append(text_elem)

        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    def _handle_list(self, node: dict[str, Any]) -> None:
        """Handle ordered and unordered lists using Word styles."""
        if not self._document:
            return

        ordered = node.get("attrs", {}).get("ordered", False)
        items = node.get("children", [])

        for item in items:
            if item.get("type") == "list_item":
                self._handle_list_item(item, ordered)

    def _handle_list_item(self, node: dict[str, Any], ordered: bool) -> None:
        """Handle a single list item using Word list styles."""
        if not self._document:
            return

        # Process list item content
        children = node.get("children", [])
        for child in children:
            if child.get("type") in ("paragraph", "block_text"):
                # Use Word's built-in list styles
                if ordered:
                    style_name = "List Number"
                    if self._list_level > 0:
                        style_name = f"List Number {self._list_level + 1}"
                else:
                    style_name = "List Bullet"
                    if self._list_level > 0:
                        style_name = f"List Bullet {self._list_level + 1}"

                # Check if style exists, fallback to base style
                try:
                    paragraph = self._document.add_paragraph(style=style_name)
                except KeyError:
                    # Fallback to base style if numbered style doesn't exist
                    paragraph = self._document.add_paragraph(
                        style="List Number" if ordered else "List Bullet"
                    )
                    # Manually adjust indentation for nested lists
                    if self._list_level > 0:
                        paragraph.paragraph_format.left_indent = Inches(0.5 * (self._list_level + 1))

                self._process_inline_content(paragraph, child.get("children", []))

            elif child.get("type") == "list":
                # Nested list
                self._list_level += 1
                self._handle_list(child)
                self._list_level -= 1

    def _handle_table(self, node: dict[str, Any]) -> None:
        """Handle tables."""
        if not self._document:
            return

        children = node.get("children", [])
        if not children:
            return

        # Get table head and body
        head = None
        body = None
        for child in children:
            if child.get("type") == "table_head":
                head = child
            elif child.get("type") == "table_body":
                body = child

        if not head:
            return

        # Get header cells (directly under table_head, not wrapped in row)
        header_cells = head.get("children", [])
        if not header_cells:
            return

        num_cols = len(header_cells)

        # Extract alignment info from header cells
        alignments = []
        for cell_node in header_cells:
            align = cell_node.get("attrs", {}).get("align")
            alignments.append(align)

        # Count total rows
        body_rows = body.get("children", []) if body else []
        num_rows = 1 + len(body_rows)  # 1 for header

        # Create table
        table = self._document.add_table(rows=num_rows, cols=num_cols)
        table.style = "Table Grid"

        # Fill header row
        for col_idx, cell_node in enumerate(header_cells):
            cell = table.rows[0].cells[col_idx]
            cell_text = MarkdownParser.get_node_text(cell_node)
            cell.text = cell_text
            # Make header bold and apply alignment
            for paragraph in cell.paragraphs:
                self._apply_table_cell_alignment(paragraph, alignments[col_idx])
                for run in paragraph.runs:
                    run.font.bold = True

        # Fill body rows
        for row_idx, row_node in enumerate(body_rows, start=1):
            row_cells = row_node.get("children", [])
            for col_idx, cell_node in enumerate(row_cells):
                if col_idx < num_cols:
                    cell = table.rows[row_idx].cells[col_idx]
                    cell_text = MarkdownParser.get_node_text(cell_node)
                    cell.text = cell_text
                    # Apply alignment
                    for paragraph in cell.paragraphs:
                        self._apply_table_cell_alignment(paragraph, alignments[col_idx])

        # Add spacing after table
        self._document.add_paragraph()

    def _apply_table_cell_alignment(self, paragraph, align: Optional[str]) -> None:
        """Apply alignment to a table cell paragraph."""
        if align == "left":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif align == "center":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # If None or unknown, leave default (left)

    def _handle_code_block(self, node: dict[str, Any]) -> None:
        """Handle fenced code blocks."""
        if not self._document:
            return

        code = node.get("raw", "").rstrip("\n")

        paragraph = self._document.add_paragraph()
        run = paragraph.add_run(code)
        run.font.name = self.branding.code_font.name
        run.font.size = self.branding.code_font.size
        if self.branding.code_font.color:
            run.font.color.rgb = self.branding.code_font.color

        # Add background shading (if configured)
        if self.branding.code_background_color:
            self._add_shading(paragraph, self.branding.code_background_color)

        # Set paragraph formatting for code block
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)

    def _add_shading(self, paragraph, color: RGBColor) -> None:
        """Add background shading to a paragraph."""
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
        paragraph._p.get_or_add_pPr().append(shd)

    def _handle_block_quote(self, node: dict[str, Any]) -> None:
        """Handle block quotes."""
        if not self._document:
            return

        children = node.get("children", [])

        for child in children:
            if child.get("type") == "paragraph":
                # Use Word's built-in Intense Quote style for better visual distinction
                paragraph = self._document.add_paragraph(style="Intense Quote")
                self._process_inline_content(paragraph, child.get("children", []))

    def _handle_thematic_break(self, node: dict[str, Any]) -> None:
        """Handle horizontal rules (thematic breaks)."""
        if not self._document:
            return

        paragraph = self._document.add_paragraph()
        paragraph.add_run("_" * 50)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(12)

    def _handle_newline(self, node: dict[str, Any]) -> None:
        """Handle explicit newlines."""
        # Usually handled by paragraph breaks, so this is a no-op
        pass
